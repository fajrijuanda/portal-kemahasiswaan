from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Manual_Book_Portal_Kemahasiswaan_UBP.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
FILL = "E8EEF5"
LIGHT = "F4F6F9"
SKY = "EAF8FF"
GREEN = "E6F7EF"
AMBER = "FFF7E6"
RED = "FDECEC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D9E2EC", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_indent(table, dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_break(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def style_table(table, widths: list[float], header_fill: str = FILL) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    set_table_indent(table, 120)
    set_cell_margins(table)
    for row_index, row in enumerate(table.rows):
        prevent_row_break(row)
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths[col_index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
    style_table(table, widths, header_fill)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, title: str, body: str, fill: str = SKY) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_table_borders(table, color="B9DDF2")
    set_table_indent(table)
    set_cell_margins(table, 110, 150, 110, 150)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(9.8)
        run.font.color.rgb = RGBColor.from_string("334155")
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str = ""):
    return doc.add_paragraph(text)


def section_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("172033")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    return doc


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Portal Kemahasiswaan UBP | Halaman ")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r2 = p.add_run()
    r2._r.append(fld_begin)
    r2._r.append(instr)
    r2._r.append(fld_end)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(MUTED)


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("MANUAL BOOK")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Portal Kemahasiswaan UBP Karawang")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Panduan operasional, alur proses, use case, dan blueprint flowchart untuk draw.io")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph()
    add_callout(
        doc,
        "Ruang Lingkup Dokumen",
        "Dokumen ini menjelaskan penggunaan portal dari sisi publik, mahasiswa, Ormawa, admin, kabag, kaprodi, warek, dan super user. Bagian flowchart disusun sebagai blueprint node dan edge agar mudah digambar ulang di draw.io.",
        SKY,
    )
    add_table(
        doc,
        ["Atribut", "Keterangan"],
        [
            ["Nama Sistem", "Portal Kemahasiswaan UBP Karawang"],
            ["Versi Manual", "1.0"],
            ["Tanggal", "30 Juni 2026"],
            ["Target Pembaca", "Pimpinan, admin kemahasiswaan, kabag, kaprodi, mahasiswa, Ormawa, dan tim pengembang"],
            ["Format", "Manual book DOCX dan blueprint flowchart untuk draw.io"],
        ],
        [1.6, 4.9],
    )
    section_break(doc)


def toc(doc: Document) -> None:
    add_h(doc, "Daftar Isi Ringkas", 1)
    rows = [
        ["1", "Gambaran umum portal"],
        ["2", "Role dan hak akses"],
        ["3", "Panduan halaman publik"],
        ["4", "Panduan login, home, dashboard, profil, dan logout"],
        ["5", "Panduan modul layanan internal"],
        ["6", "Panduan modul Ormawa, Master, Publikasi, User"],
        ["7", "Blueprint flowchart per modul untuk draw.io"],
        ["8", "Use case description per modul"],
        ["9", "Checklist QA dan troubleshooting"],
    ]
    add_table(doc, ["No.", "Bagian"], rows, [0.7, 5.8])


def overview(doc: Document) -> None:
    add_h(doc, "1. Gambaran Umum Portal", 1)
    add_p(
        doc,
        "Portal Kemahasiswaan UBP Karawang adalah aplikasi web untuk mengelola layanan kemahasiswaan secara terpusat. Sistem menyediakan halaman publik untuk informasi umum serta panel internal untuk pengajuan, verifikasi, rekap, master data, publikasi berita, karir, dan manajemen pengguna.",
    )
    add_callout(
        doc,
        "Prinsip Utama",
        "Semua pengajuan dari mahasiswa dan Ormawa masuk dengan status awal Diajukan. Admin, kabag, kaprodi, atau role terkait kemudian melakukan review sesuai kewenangannya. Data yang sudah valid akan masuk ke rekap dan dashboard.",
        GREEN,
    )
    add_h(doc, "1.1 Modul Utama", 2)
    add_table(
        doc,
        ["Modul", "Fungsi Utama", "Aktor Utama"],
        [
            ["Publik", "Landing page, profil, layanan, berita, links, detail layanan", "Pengunjung, mahasiswa, alumni, mitra"],
            ["Login/Auth", "Autentikasi, CSRF, redirect ke Home, logout confirm", "Semua user"],
            ["Home", "Katalog layanan internal dan akses cepat sesuai role", "User login"],
            ["Dashboard", "Monitoring statistik, grafik, dan rekap", "Admin, kabag, kaprodi, warek, super user"],
            ["Prestasi", "Data prestasi lomba, kategori, scope, juara, dan kuota prodi", "Mahasiswa, admin, kaprodi"],
            ["Event", "Data kegiatan mahasiswa dan pengajuan event", "Mahasiswa, admin"],
            ["Reimbursement", "Klaim biaya mahasiswa atau Ormawa dengan dokumen pendukung", "Mahasiswa, Ormawa, admin"],
            ["Beasiswa", "Jenis beasiswa, nominal, status, dan pengajuan mahasiswa", "Mahasiswa, admin"],
            ["Tracer Study", "Pendataan tracer/alumni", "Admin, alumni/mahasiswa"],
            ["Unit", "Humas, Science Center, Alumni/Pusat Karir", "Admin unit, kabag"],
            ["Ormawa", "Master Ormawa, kegiatan, proposal, reimbursement Ormawa", "Admin, Ormawa"],
            ["Master", "Prodi, semester, lomba, jenis beasiswa, kuota prestasi", "Admin, super user"],
            ["Publikasi", "Berita, loker, dan job fair", "Kabag, admin, super user"],
            ["User", "Kelola akun, role, prodi, NIM, password", "Super user"],
            ["Profil", "Update profil, password, dan penghapusan akun", "Semua user"],
        ],
        [1.35, 3.35, 1.8],
    )


def roles(doc: Document) -> None:
    add_h(doc, "2. Role dan Hak Akses", 1)
    add_table(
        doc,
        ["Role", "Warna Badge", "Hak Akses Ringkas"],
        [
            ["Super User", "Violet", "Akses penuh termasuk Management User, Master, Publikasi, Dashboard, dan seluruh layanan admin."],
            ["Admin", "Biru", "Mengelola data layanan, unit, Ormawa, master data, karir, dan dashboard operasional."],
            ["Kabag", "Rose", "Membuat dan mempublikasikan berita serta memantau layanan sesuai bagian."],
            ["Kaprodi", "Emerald", "Melihat dan memverifikasi data sesuai scope prodi, terutama prestasi dan kuota."],
            ["Warek", "Amber", "Akses monitoring dan rekap pimpinan."],
            ["Mahasiswa", "Cyan", "Mengajukan beasiswa dan prestasi/lomba melalui panel pengajuan."],
            ["Ormawa", "Orange", "Mengajukan proposal kegiatan dan reimbursement acara Ormawa."],
        ],
        [1.35, 1.15, 4.0],
    )
    add_h(doc, "2.1 Matriks Akses Modul", 2)
    add_table(
        doc,
        ["Modul", "Super User", "Admin", "Kabag", "Kaprodi", "Warek", "Mahasiswa", "Ormawa"],
        [
            ["Home", "Ya", "Ya", "Ya", "Ya", "Ya", "Ya", "Ya"],
            ["Dashboard", "Ya", "Ya", "Ya", "Ya", "Ya", "Tidak", "Tidak"],
            ["Prestasi/Event/Reimburse/Beasiswa/Tracer", "Ya", "Ya", "Lihat sesuai kebutuhan", "Scope prodi", "Lihat", "Pengajuan", "Tidak"],
            ["Unit", "Ya", "Ya", "Ya", "Lihat", "Lihat", "Tidak", "Tidak"],
            ["Ormawa Admin", "Ya", "Ya", "Lihat", "Tidak", "Lihat", "Tidak", "Tidak"],
            ["Panel Ormawa", "Tidak", "Tidak", "Tidak", "Tidak", "Tidak", "Tidak", "Ya"],
            ["Master", "Ya", "Ya", "Tidak", "Tidak", "Tidak", "Tidak", "Tidak"],
            ["Publikasi", "Ya", "Ya", "Berita", "Tidak", "Lihat", "Tidak", "Tidak"],
            ["Management User", "Ya", "Tidak", "Tidak", "Tidak", "Tidak", "Tidak", "Tidak"],
        ],
        [1.35, 0.75, 0.7, 0.7, 0.72, 0.65, 0.83, 0.8],
    )


def public_manual(doc: Document) -> None:
    add_h(doc, "3. Panduan Halaman Publik", 1)
    add_p(doc, "Route `/` menampilkan landing page publik. Navbar publik berpindah halaman ke Beranda, Profil, Layanan, Berita, Links, dan Login.")
    add_h(doc, "3.1 Halaman Publik", 2)
    add_table(
        doc,
        ["Route", "Nama Halaman", "Isi Utama"],
        [
            ["/", "Beranda", "Hero portal, statistik ringkas, layanan utama, berita terbaru, karir, FAQ, dan tombol login."],
            ["/profil", "Profil", "Profil layanan kemahasiswaan, fokus layanan, transparansi proses, publikasi resmi, dan rekap pimpinan."],
            ["/layanan", "Layanan", "Katalog layanan publik dengan card icon + title yang dapat diklik."],
            ["/layanan/{slug}", "Detail Layanan", "Ringkasan layanan, informasi pengajuan, dan link login panel internal."],
            ["/berita", "Berita", "Daftar berita published dari admin/kabag."],
            ["/berita/{slug}", "Detail Berita", "Isi berita published dengan cover, tanggal, ringkasan, dan konten."],
            ["/links", "Links", "Tautan cepat publik dan daftar loker/job fair published."],
        ],
        [1.35, 1.65, 3.5],
    )
    add_h(doc, "3.2 Layanan Publik", 2)
    add_table(
        doc,
        ["Layanan", "Slug", "Ringkasan"],
        [
            ["Prestasi Mahasiswa", "prestasi-mahasiswa", "Prestasi lomba, kategori event, scope, juara, dan kuota prestasi prodi."],
            ["Event Mahasiswa", "event-mahasiswa", "Pengajuan kegiatan, dokumentasi, dan status review kegiatan mahasiswa."],
            ["Reimbursement", "reimbursement", "Pengajuan klaim biaya dengan foto, surat tugas, sertifikat, dan link penyelenggara."],
            ["Beasiswa", "beasiswa", "Pengajuan KIP, Kacer, Tahfidz, dan jenis beasiswa lainnya."],
            ["Ormawa", "ormawa", "Profil organisasi, proposal kegiatan, dan reimbursement acara Ormawa."],
            ["Tracer Study", "tracer-study", "Pengumpulan data alumni untuk evaluasi dan akreditasi."],
            ["Berita", "berita", "Berita resmi kemahasiswaan yang sudah published."],
            ["Karir", "karir", "Lowongan kerja dan job fair untuk mahasiswa serta alumni."],
        ],
        [1.6, 1.35, 3.55],
    )
    add_h(doc, "3.3 Langkah Pengunjung Publik", 2)
    add_numbered(
        doc,
        [
            "Buka route `/` untuk melihat beranda publik.",
            "Pilih menu navbar: Profil, Layanan, Berita, atau Links.",
            "Pada halaman Layanan, klik card layanan untuk membuka detail layanan.",
            "Pada halaman Berita, klik salah satu berita untuk membuka detail.",
            "Jika ingin melakukan pengajuan atau pengelolaan data, klik Login dan masuk memakai akun portal.",
        ],
    )


def auth_manual(doc: Document) -> None:
    add_h(doc, "4. Login, Home, Dashboard, Profil, dan Logout", 1)
    add_h(doc, "4.1 Login", 2)
    add_numbered(
        doc,
        [
            "Buka `/login` atau klik tombol Login dari halaman publik.",
            "Isi email dan password.",
            "Klik Login.",
            "Jika autentikasi berhasil, user diarahkan ke `/home`.",
            "Jika gagal, sistem menampilkan validasi error.",
        ],
    )
    add_callout(doc, "Catatan Keamanan", "Form login harus dikirim melalui HTTPS agar browser tidak menampilkan peringatan form tidak aman dan untuk mencegah error CSRF 419.", AMBER)
    add_h(doc, "4.2 Home", 2)
    add_p(doc, "Home adalah halaman awal setelah login. Katalog layanan di Home mengikuti menu baru dan role user. Home berbeda dari Dashboard; Dashboard dipakai untuk rekap dan monitoring.")
    add_h(doc, "4.3 Dashboard", 2)
    add_p(doc, "Dashboard menampilkan rekap layanan, grafik, dan statistik. Menu ini ditujukan untuk admin, kabag, kaprodi, warek, dan super user.")
    add_h(doc, "4.4 Profil", 2)
    add_p(doc, "Halaman Profil menggunakan template portal yang sama dengan halaman lain. User dapat memperbarui nama, email, password, dan menghapus akun dengan konfirmasi password.")
    add_h(doc, "4.5 Logout", 2)
    add_p(doc, "Logout dilakukan dari dropdown profil. Sistem menampilkan modal konfirmasi sebelum sesi diakhiri.")


def module_manual(doc: Document) -> None:
    add_h(doc, "5. Panduan Modul Layanan Internal", 1)
    modules = [
        ("Prestasi", "Mencatat prestasi lomba mahasiswa, competition_id, kategori event, scope, juara, prodi, semester, dan status.", ["Buka menu Prestasi.", "Pilih tabel Prestasi Mahasiswa.", "Gunakan search/filter semester/prodi.", "Klik Tambah Data atau Edit.", "Set status sesuai hasil review. Jika Terverifikasi, kuota prodi bertambah."]),
        ("Event", "Mencatat kegiatan/event mahasiswa secara terpisah dari reimbursement.", ["Buka menu Event.", "Pilih mini sidebar Event/Kegiatan.", "Input data kegiatan, mahasiswa/pengaju, prodi, semester, dan status.", "Review status sampai selesai."]),
        ("Reimbursement", "Mengelola klaim biaya mahasiswa atau Ormawa dengan syarat dokumen.", ["Buka menu Event.", "Pilih mini sidebar Reimbursement.", "Pastikan foto, surat tugas, sertifikat, dan link penyelenggara tersedia.", "Review nominal, status, dan catatan."]),
        ("Beasiswa", "Menampilkan nama mahasiswa, jenis beasiswa, nominal, prodi, dan status.", ["Buka menu Beasiswa.", "Pilih tabel data beasiswa.", "Gunakan filter/search.", "Review pengajuan mahasiswa.", "Update status Diajukan, Diproses, Disetujui, atau Ditolak."]),
        ("Tracer Study", "Mengelola data tracer/alumni dan status kelengkapan input.", ["Buka menu Tracer.", "Pilih tabel Tracer Study.", "Cari data mahasiswa/alumni.", "Update data sesuai hasil input atau follow up."]),
        ("Unit", "Mengelola data unit Humas Marketing, Science Center, dan Alumni/Pusat Karir.", ["Buka menu Unit.", "Pilih mini sidebar unit.", "Input aktivitas unit.", "Gunakan filter dan tabel untuk monitoring."]),
    ]
    for title, desc, steps in modules:
        add_h(doc, title, 2)
        add_p(doc, desc)
        add_numbered(doc, steps)


def admin_manual(doc: Document) -> None:
    add_h(doc, "6. Modul Ormawa, Master, Publikasi, User", 1)
    add_h(doc, "6.1 Ormawa Admin", 2)
    add_p(doc, "Menu Ormawa menyatukan Data Ormawa, Kegiatan, Proposal, dan Reimbursement Ormawa. Mini sidebar digunakan agar CRUD yang satu lingkup tetap berada dalam satu halaman keluarga.")
    add_bullets(doc, ["Data Ormawa: master profil organisasi, jenis, pembina/PIC, kontak, deskripsi, status, dan akun Ormawa.", "Kegiatan: gambaran kegiatan Ormawa atau unit pengembangan Ormawa.", "Proposal: pengajuan proposal dari akun Ormawa.", "Reimbursement: klaim biaya acara yang terhubung ke ormawa_id."])
    add_h(doc, "6.2 Panel Ormawa", 2)
    add_p(doc, "Role Ormawa memiliki panel sendiri untuk mengajukan proposal kegiatan dan reimbursement acara. Status awal semua pengajuan adalah Diajukan.")
    add_h(doc, "6.3 Master Data", 2)
    add_p(doc, "Master data menyediakan Prodi, Semester, Lomba, Jenis Beasiswa, dan Kuota Prestasi. Tabel master dilengkapi search dan filter sesuai konteks.")
    add_table(
        doc,
        ["Master", "Fungsi", "Filter"],
        [
            ["Prodi", "Nama, kode, fakultas", "Search dan fakultas"],
            ["Semester", "Nama, tahun akademik, periode, status aktif", "Search, periode, status"],
            ["Lomba", "Master 23 nama lomba dan status aktif", "Search, status"],
            ["Jenis Beasiswa", "KIP, Kacer, Tahfidz, Lainnya, dan tambahan admin", "Search, status"],
            ["Kuota Prestasi", "Slot prestasi per prodi dan semester", "Search, semester, prodi"],
        ],
        [1.35, 3.1, 2.05],
    )
    add_h(doc, "6.4 Publikasi", 2)
    add_p(doc, "Publikasi terdiri dari Berita dan Karir. Berita dibuat oleh kabag/admin dan tampil di halaman publik jika status Published. Karir berisi Loker dan Job Fair untuk halaman Links.")
    add_h(doc, "6.5 Management User", 2)
    add_p(doc, "Management User hanya untuk Super User. Fitur ini mengelola akun, role, prodi, NIM, password, dan penghapusan akun. Badge role memiliki icon dan warna berbeda untuk memudahkan identifikasi.")


def flowchart_intro(doc: Document) -> None:
    section_break(doc)
    add_h(doc, "7. Blueprint Flowchart untuk draw.io", 1)
    add_p(doc, "Bagian ini menjelaskan node dan edge flowchart. Gunakan terminator untuk Start/End, process untuk aksi, decision untuk percabangan, document untuk dokumen/file, dan database untuk penyimpanan data.")
    add_table(
        doc,
        ["Simbol draw.io", "Dipakai Untuk", "Contoh Label"],
        [
            ["Terminator", "Awal/akhir alur", "Start, End"],
            ["Process", "Aktivitas user/sistem", "Isi form login, Simpan pengajuan"],
            ["Decision", "Percabangan ya/tidak", "Data valid? Role sesuai?"],
            ["Document", "File pendukung", "Upload sertifikat"],
            ["Database", "Penyimpanan data", "Tabel beasiswas"],
            ["Connector", "Penghubung antar proses", "Ke review admin"],
        ],
        [1.55, 2.35, 2.6],
    )


FLOWCHARTS = [
    (
        "Login dan Redirect Home",
        [
            ["F1-01", "Start", "User membuka /login", "F1-02"],
            ["F1-02", "Process", "User mengisi email dan password", "F1-03"],
            ["F1-03", "Decision", "Credential valid?", "Ya: F1-04; Tidak: F1-08"],
            ["F1-04", "Decision", "CSRF/session valid?", "Ya: F1-05; Tidak: F1-09"],
            ["F1-05", "Process", "Sistem membuat session login", "F1-06"],
            ["F1-06", "Process", "Redirect ke /home", "F1-07"],
            ["F1-07", "End", "Home tampil sesuai role", "-"],
            ["F1-08", "Process", "Tampilkan error login", "F1-02"],
            ["F1-09", "Process", "Tampilkan error 419/session expired", "F1-02"],
        ],
    ),
    (
        "Halaman Publik dan Detail Layanan",
        [
            ["F2-01", "Start", "Pengunjung membuka /", "F2-02"],
            ["F2-02", "Process", "Sistem menampilkan landing public", "F2-03"],
            ["F2-03", "Decision", "Klik menu navbar?", "Profil/Layanan/Berita/Links/Login"],
            ["F2-04", "Process", "Jika Layanan, tampilkan katalog card", "F2-05"],
            ["F2-05", "Decision", "Klik card layanan?", "Ya: F2-06; Tidak: F2-07"],
            ["F2-06", "Process", "Buka /layanan/{slug}", "F2-07"],
            ["F2-07", "Decision", "Perlu login?", "Ya: F2-08; Tidak: End"],
            ["F2-08", "Process", "Klik Login", "F1-01"],
        ],
    ),
    (
        "Pengajuan Beasiswa Mahasiswa",
        [
            ["F3-01", "Start", "Mahasiswa login", "F3-02"],
            ["F3-02", "Process", "Buka menu Pengajuan", "F3-03"],
            ["F3-03", "Process", "Isi form beasiswa: jenis, nominal, catatan", "F3-04"],
            ["F3-04", "Decision", "Data valid?", "Ya: F3-05; Tidak: F3-08"],
            ["F3-05", "Database", "Simpan ke beasiswas status Diajukan", "F3-06"],
            ["F3-06", "Process", "Admin review di menu Beasiswa", "F3-07"],
            ["F3-07", "Decision", "Disetujui?", "Ya: status Disetujui; Tidak: status Ditolak/Diproses"],
            ["F3-08", "Process", "Tampilkan validasi", "F3-03"],
        ],
    ),
    (
        "Pengajuan Prestasi/Lomba Mahasiswa",
        [
            ["F4-01", "Start", "Mahasiswa login", "F4-02"],
            ["F4-02", "Process", "Pilih pengajuan prestasi", "F4-03"],
            ["F4-03", "Process", "Isi lomba, kategori event, scope, juara, prodi", "F4-04"],
            ["F4-04", "Decision", "Data valid?", "Ya: F4-05; Tidak: F4-09"],
            ["F4-05", "Database", "Simpan prestasi status Diajukan", "F4-06"],
            ["F4-06", "Process", "Admin/Kaprodi review", "F4-07"],
            ["F4-07", "Decision", "Terverifikasi?", "Ya: F4-08; Tidak: status Ditolak/Diproses"],
            ["F4-08", "Database", "Tambah terpakai pada kuota prodi", "End"],
            ["F4-09", "Process", "Tampilkan validasi", "F4-03"],
        ],
    ),
    (
        "Event dan Reimbursement Mahasiswa",
        [
            ["F5-01", "Start", "Admin membuka menu Event", "F5-02"],
            ["F5-02", "Decision", "Pilih mini sidebar?", "Event: F5-03; Reimbursement: F5-06"],
            ["F5-03", "Process", "Input/update data event", "F5-04"],
            ["F5-04", "Database", "Simpan ke events module event", "F5-05"],
            ["F5-05", "End", "Event tampil di tabel"],
            ["F5-06", "Process", "Input reimbursement dan dokumen", "F5-07"],
            ["F5-07", "Decision", "File dan URL valid?", "Ya: F5-08; Tidak: F5-10"],
            ["F5-08", "Database", "Simpan reimbursement", "F5-09"],
            ["F5-09", "End", "Reimbursement siap review"],
            ["F5-10", "Process", "Tampilkan validasi", "F5-06"],
        ],
    ),
    (
        "Ormawa Proposal dan Reimbursement",
        [
            ["F6-01", "Start", "Akun Ormawa login", "F6-02"],
            ["F6-02", "Process", "Buka Panel Ormawa", "F6-03"],
            ["F6-03", "Decision", "Pilih aksi?", "Proposal: F6-04; Reimbursement: F6-07"],
            ["F6-04", "Process", "Isi judul dan upload proposal", "F6-05"],
            ["F6-05", "Database", "Simpan proposal status Diajukan", "F6-06"],
            ["F6-06", "Process", "Admin review di Ormawa Admin", "End"],
            ["F6-07", "Process", "Isi reimbursement acara dan dokumen", "F6-08"],
            ["F6-08", "Database", "Simpan events dengan ormawa_id", "F6-09"],
            ["F6-09", "Process", "Admin review reimbursement Ormawa", "End"],
        ],
    ),
    (
        "Publikasi Berita dan Karir",
        [
            ["F7-01", "Start", "Kabag/Admin membuka Publikasi", "F7-02"],
            ["F7-02", "Decision", "Pilih section?", "Berita: F7-03; Karir: F7-07"],
            ["F7-03", "Process", "Isi judul, ringkasan, konten, cover, status", "F7-04"],
            ["F7-04", "Decision", "Status Published?", "Ya: F7-05; Tidak: F7-06"],
            ["F7-05", "Database", "Simpan published_at dan tampil di /berita", "End"],
            ["F7-06", "Database", "Simpan Draft, tidak tampil publik", "End"],
            ["F7-07", "Process", "Isi loker/job fair, perusahaan, deadline, URL", "F7-08"],
            ["F7-08", "Decision", "Status Published?", "Ya tampil di /links; Tidak Draft"],
        ],
    ),
    (
        "Management User",
        [
            ["F8-01", "Start", "Super User membuka Management User", "F8-02"],
            ["F8-02", "Decision", "Aksi?", "Tambah/Edit/Hapus/Filter"],
            ["F8-03", "Process", "Tambah/Edit: isi nama, email, role, prodi, NIM, password", "F8-04"],
            ["F8-04", "Decision", "Valid?", "Ya: F8-05; Tidak: F8-07"],
            ["F8-05", "Database", "Simpan user dan sync role", "F8-06"],
            ["F8-06", "End", "User tampil dengan role badge", "-"],
            ["F8-07", "Process", "Tampilkan validasi", "F8-03"],
        ],
    ),
    (
        "Master Data dan Kuota Prestasi",
        [
            ["F9-01", "Start", "Admin membuka Master", "F9-02"],
            ["F9-02", "Decision", "Pilih section?", "Prodi/Semester/Lomba/Jenis Beasiswa/Kuota"],
            ["F9-03", "Process", "Gunakan search/filter jika perlu", "F9-04"],
            ["F9-04", "Decision", "Tambah/Edit/Hapus?", "Ya: F9-05; Tidak: End"],
            ["F9-05", "Process", "Isi form master", "F9-06"],
            ["F9-06", "Decision", "Valid?", "Ya: F9-07; Tidak: F9-08"],
            ["F9-07", "Database", "Simpan master data", "End"],
            ["F9-08", "Process", "Tampilkan validasi", "F9-05"],
        ],
    ),
    (
        "Profil dan Logout",
        [
            ["F10-01", "Start", "User membuka dropdown profil", "F10-02"],
            ["F10-02", "Decision", "Pilih Profil atau Logout?", "Profil: F10-03; Logout: F10-07"],
            ["F10-03", "Process", "Buka /profile", "F10-04"],
            ["F10-04", "Decision", "Update profil/password/hapus akun?", "Sesuai pilihan"],
            ["F10-05", "Process", "Submit form dan validasi", "F10-06"],
            ["F10-06", "End", "Data akun diperbarui atau akun dihapus", "-"],
            ["F10-07", "Process", "Tampilkan modal konfirmasi logout", "F10-08"],
            ["F10-08", "Decision", "Konfirmasi?", "Ya: destroy session; Tidak: batal"],
        ],
    ),
]


def flowcharts(doc: Document) -> None:
    for title, rows in FLOWCHARTS:
        add_h(doc, f"7.x {title}", 2)
        add_table(doc, ["Node ID", "Simbol", "Label", "Edge/Next"], rows, [0.8, 1.0, 3.2, 1.5])


USE_CASES = [
    ["UC-01", "Login", "Semua user", "User memiliki akun aktif", "User mengisi email/password dan masuk ke Home", "Session aktif dan Home tampil sesuai role", "Credential salah atau CSRF/session expired"],
    ["UC-02", "Melihat halaman publik", "Pengunjung", "Tidak perlu login", "Pengunjung membuka Beranda/Profil/Layanan/Berita/Links", "Informasi publik tampil", "Konten berita/karir kosong jika belum published"],
    ["UC-03", "Membuka detail layanan", "Pengunjung", "Halaman Layanan tersedia", "Klik card layanan", "Detail layanan tampil dengan tombol Login", "Slug tidak valid menghasilkan 404"],
    ["UC-04", "Submit beasiswa", "Mahasiswa", "Login role mahasiswa", "Isi form jenis beasiswa, nominal, catatan", "Data masuk status Diajukan", "Data tidak valid"],
    ["UC-05", "Submit prestasi/lomba", "Mahasiswa", "Login role mahasiswa dan prodi tersedia", "Isi data lomba, kategori, scope, juara", "Prestasi status Diajukan", "Master lomba belum tersedia"],
    ["UC-06", "Review prestasi", "Admin/Kaprodi", "Ada data Diajukan", "Edit status menjadi Terverifikasi/Ditolak", "Status berubah dan kuota bertambah jika Terverifikasi", "Kuota tidak ditemukan atau data tidak valid"],
    ["UC-07", "Kelola event", "Admin", "Login admin", "Tambah/edit/hapus event", "Event tersimpan dan tampil di tabel", "Validasi gagal"],
    ["UC-08", "Kelola reimbursement", "Admin/Ormawa", "Dokumen tersedia", "Input nominal, dokumen, link penyelenggara", "Reimbursement tersimpan", "File tidak valid atau URL tidak valid"],
    ["UC-09", "Kelola Ormawa", "Admin", "Login admin", "Tambah/edit master Ormawa dan lihat overview", "Data Ormawa dan relasi akun tersimpan", "Akun belum dipilih"],
    ["UC-10", "Submit proposal Ormawa", "Ormawa", "Login akun Ormawa", "Upload proposal kegiatan", "Proposal status Diajukan", "File tidak valid"],
    ["UC-11", "Kelola master data", "Admin/Super User", "Login dan akses Master", "Tambah/edit/hapus prodi, semester, lomba, jenis beasiswa, kuota", "Master tersedia di dropdown/filter", "Data duplikat atau relasi masih dipakai"],
    ["UC-12", "Publikasi berita", "Kabag/Admin", "Login dan akses Publikasi", "Isi berita dan set Published", "Berita tampil di /berita", "Draft tidak tampil publik"],
    ["UC-13", "Publikasi karir", "Admin/Super User", "Login dan akses Publikasi Karir", "Input loker/job fair dan set Published", "Konten tampil di /links", "URL eksternal invalid"],
    ["UC-14", "Management user", "Super User", "Login super user", "Tambah/edit/hapus akun dan role", "User tersimpan dengan badge role", "Email duplikat atau role invalid"],
    ["UC-15", "Update profil", "Semua user", "Login", "Edit nama/email atau password", "Profil/password tersimpan", "Password saat ini salah"],
    ["UC-16", "Logout", "Semua user", "Login", "Klik Logout dan konfirmasi modal", "Session berakhir", "User membatalkan modal"],
]


def use_cases(doc: Document) -> None:
    section_break(doc)
    add_h(doc, "8. Use Case Description", 1)
    add_p(doc, "Tabel berikut dapat digunakan sebagai bahan penulisan use case diagram dan use case specification.")
    add_table(
        doc,
        ["ID", "Use Case", "Aktor", "Precondition", "Main Flow", "Postcondition", "Exception"],
        USE_CASES,
        [0.55, 1.1, 0.9, 1.2, 1.45, 1.2, 1.1],
    )


def troubleshooting(doc: Document) -> None:
    add_h(doc, "9. Checklist QA dan Troubleshooting", 1)
    add_h(doc, "9.1 Checklist Operasional", 2)
    add_bullets(
        doc,
        [
            "Pastikan aplikasi diakses melalui HTTPS, terutama untuk form login dan upload.",
            "Pastikan user memiliki role yang benar sebelum menguji menu.",
            "Pastikan master prodi, semester, lomba, dan jenis beasiswa tersedia sebelum pengajuan.",
            "Pastikan file reimbursement berformat pdf, jpg, jpeg, atau png dan berukuran maksimal 2 MB.",
            "Pastikan berita, loker, dan job fair berstatus Published agar tampil di halaman publik.",
            "Pastikan route lama tetap redirect dan tidak menghasilkan 404.",
        ],
    )
    add_h(doc, "9.2 Masalah Umum", 2)
    add_table(
        doc,
        ["Masalah", "Kemungkinan Penyebab", "Solusi"],
        [
            ["Error 419 saat login/submit", "CSRF token expired, session hilang, atau form dikirim tidak aman", "Refresh halaman, pastikan HTTPS, login ulang, dan cek konfigurasi session."],
            ["Form not secure di browser", "Form submit ke endpoint HTTP atau mixed content", "Pastikan seluruh URL memakai HTTPS dan APP_URL sesuai domain production."],
            ["Menu tidak muncul", "Role user tidak sesuai", "Cek Management User dan sync role yang benar."],
            ["Data tidak tampil di publik", "Status masih Draft", "Ubah status ke Published."],
            ["Filter master tidak menemukan data", "Keyword/filter terlalu spesifik", "Reset filter atau gunakan keyword lebih umum."],
            ["Upload reimbursement gagal", "File type/size tidak sesuai atau URL penyelenggara invalid", "Gunakan pdf/jpg/jpeg/png maksimal 2 MB dan URL valid."],
        ],
        [1.6, 2.4, 2.5],
    )
    add_h(doc, "9.3 Catatan untuk draw.io", 2)
    add_bullets(
        doc,
        [
            "Gunakan satu warna per aktor utama: Mahasiswa, Ormawa, Admin/Kabag, dan Sistem.",
            "Pisahkan swimlane jika flowchart melibatkan lebih dari satu aktor.",
            "Gunakan database shape untuk tabel seperti beasiswas, prestasis, events, ormawa_proposals, press_releases, career_posts, dan users.",
            "Gunakan decision diamond untuk validasi, status Published, role access, dan approve/reject.",
            "Label edge sebaiknya singkat: Ya, Tidak, Draft, Published, Disetujui, Ditolak.",
        ],
    )


def final_notes(doc: Document) -> None:
    section_break(doc)
    add_h(doc, "Lampiran A. Ringkasan Route Penting", 1)
    add_table(
        doc,
        ["Area", "Route", "Keterangan"],
        [
            ["Publik", "/", "Beranda publik"],
            ["Publik", "/profil", "Profil layanan publik"],
            ["Publik", "/layanan", "Katalog layanan"],
            ["Publik", "/layanan/{slug}", "Detail layanan"],
            ["Publik", "/berita", "Daftar berita published"],
            ["Publik", "/berita/{slug}", "Detail berita"],
            ["Publik", "/links", "Tautan dan karir publik"],
            ["Auth", "/login", "Login user"],
            ["Internal", "/home", "Home setelah login"],
            ["Internal", "/dashboard", "Dashboard rekap"],
            ["Layanan", "/prestasi, /event, /beasiswa, /tracer", "Overview modul layanan"],
            ["Layanan", "/prestasi/mahasiswa, /event/kegiatan, /event/reimburse", "Tabel spesifik"],
            ["Unit", "/unit/{unit}", "Humas, Science, Alumni/Pusat Karir"],
            ["Ormawa", "/ormawa/{section}", "Data Ormawa, kegiatan, proposal, reimbursement"],
            ["Master", "/master/{section}", "Prodi, semester, competitions, scholarship-types, quotas"],
            ["Publikasi", "/publikasi/berita, /publikasi/careers", "Berita, loker, job fair"],
            ["User", "/management-user", "Manajemen user khusus super user"],
            ["Profil", "/profile", "Profil akun"],
        ],
        [1.1, 2.1, 3.3],
    )
    add_h(doc, "Lampiran B. Status Data", 1)
    add_table(
        doc,
        ["Status", "Makna", "Dipakai Pada"],
        [
            ["Diajukan", "Data baru masuk dan menunggu review", "Pengajuan mahasiswa, Ormawa, reimbursement"],
            ["Diproses", "Data sedang diperiksa atau membutuhkan tindak lanjut", "Beasiswa, reimbursement, prestasi"],
            ["Terverifikasi", "Data prestasi sudah valid dan dapat mempengaruhi kuota", "Prestasi"],
            ["Disetujui", "Pengajuan diterima", "Beasiswa, reimbursement, proposal"],
            ["Ditolak", "Pengajuan tidak diterima", "Seluruh pengajuan"],
            ["Draft", "Konten belum tampil di publik", "Berita, karir"],
            ["Published", "Konten tampil di halaman publik", "Berita, karir"],
            ["Aktif/Nonaktif", "Master dapat/tidak dapat digunakan", "Semester, lomba, jenis beasiswa, Ormawa"],
        ],
        [1.25, 3.0, 2.25],
    )


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    for section in doc.sections:
        add_page_number(section)
    cover(doc)
    toc(doc)
    overview(doc)
    roles(doc)
    public_manual(doc)
    auth_manual(doc)
    module_manual(doc)
    admin_manual(doc)
    flowchart_intro(doc)
    flowcharts(doc)
    use_cases(doc)
    troubleshooting(doc)
    final_notes(doc)
    for section in doc.sections:
        add_page_number(section)
    doc.save(OUT)


if __name__ == "__main__":
    build()
